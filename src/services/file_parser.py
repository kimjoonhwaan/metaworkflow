"""File parsing service for extracting text from various file formats"""

import io
import os
from typing import Dict, Any, Optional, List
from pathlib import Path
import mimetypes

# File processing libraries
import PyPDF2
from docx import Document
import openpyxl
from PIL import Image
import pytesseract
try:
    import magic
except ImportError:
    magic = None

from ..utils.logger import get_logger

logger = get_logger(__name__)


class FileParser:
    """Service for parsing various file formats and extracting text content"""
    
    def __init__(self):
        # Supported file types
        self.supported_types = {
            'application/pdf': self._parse_pdf,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': self._parse_docx,
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': self._parse_xlsx,
            'application/vnd.ms-excel': self._parse_xlsx,
            'text/plain': self._parse_text,
            'text/csv': self._parse_csv,
            'image/jpeg': self._parse_image,
            'image/png': self._parse_image,
            'image/gif': self._parse_image,
            'image/bmp': self._parse_image,
            'image/tiff': self._parse_image,
        }
        
        # File extensions mapping
        self.extension_mapping = {
            '.pdf': 'application/pdf',
            '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            '.doc': 'application/msword',
            '.xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            '.xls': 'application/vnd.ms-excel',
            '.txt': 'text/plain',
            '.csv': 'text/csv',
            '.jpg': 'image/jpeg',
            '.jpeg': 'image/jpeg',
            '.png': 'image/png',
            '.gif': 'image/gif',
            '.bmp': 'image/bmp',
            '.tiff': 'image/tiff',
            '.tif': 'image/tiff',
        }
    
    def get_file_info(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Get file information including type, size, and metadata"""
        try:
            # Get file extension
            file_ext = Path(filename).suffix.lower()
            
            # Try to detect MIME type using magic if available
            if magic:
                try:
                    mime_type = magic.from_buffer(file_content, mime=True)
                except Exception:
                    mime_type = self.extension_mapping.get(file_ext, 'unknown')
            else:
                # Fallback to extension-based detection
                mime_type = self.extension_mapping.get(file_ext, 'unknown')
            
            return {
                'filename': filename,
                'mime_type': mime_type,
                'size': len(file_content),
                'extension': file_ext,
                'is_supported': mime_type in self.supported_types
            }
        except Exception as e:
            logger.error(f"Failed to get file info: {e}")
            return {
                'filename': filename,
                'mime_type': 'unknown',
                'size': len(file_content),
                'extension': Path(filename).suffix.lower(),
                'is_supported': False,
                'error': str(e)
            }
    
    def parse_file(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Parse file and extract text content"""
        try:
            file_info = self.get_file_info(file_content, filename)
            
            if not file_info['is_supported']:
                return {
                    'success': False,
                    'error': f"Unsupported file type: {file_info['mime_type']}",
                    'file_info': file_info
                }
            
            # Parse content based on file type
            parser_func = self.supported_types[file_info['mime_type']]
            content = parser_func(file_content)
            
            return {
                'success': True,
                'content': content,
                'file_info': file_info,
                'metadata': self._extract_metadata(file_content, file_info)
            }
            
        except Exception as e:
            logger.error(f"Failed to parse file {filename}: {e}")
            return {
                'success': False,
                'error': str(e),
                'file_info': self.get_file_info(file_content, filename)
            }
    
    def _parse_pdf(self, file_content: bytes) -> str:
        """Parse PDF file and extract text"""
        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            text = ""
            
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                page_text = page.extract_text()
                if page_text:
                    text += f"\n--- Page {page_num + 1} ---\n"
                    text += page_text
            
            return text.strip()
        except Exception as e:
            logger.error(f"Failed to parse PDF: {e}")
            raise
    
    def _parse_docx(self, file_content: bytes) -> str:
        """Parse DOCX file and extract text"""
        try:
            doc = Document(io.BytesIO(file_content))
            text = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text += " | ".join(row_text) + "\n"
            
            return text.strip()
        except Exception as e:
            logger.error(f"Failed to parse DOCX: {e}")
            raise
    
    def _parse_xlsx(self, file_content: bytes) -> str:
        """Parse Excel file and extract text"""
        try:
            workbook = openpyxl.load_workbook(io.BytesIO(file_content))
            text = ""
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                text += f"\n--- Sheet: {sheet_name} ---\n"
                
                for row in sheet.iter_rows():
                    row_data = []
                    for cell in row:
                        if cell.value is not None:
                            row_data.append(str(cell.value))
                    
                    if row_data:
                        text += " | ".join(row_data) + "\n"
            
            return text.strip()
        except Exception as e:
            logger.error(f"Failed to parse Excel: {e}")
            raise
    
    def _parse_text(self, file_content: bytes) -> str:
        """Parse plain text file"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'cp949', 'euc-kr', 'latin-1']
            
            for encoding in encodings:
                try:
                    return file_content.decode(encoding)
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, use utf-8 with error handling
            return file_content.decode('utf-8', errors='replace')
        except Exception as e:
            logger.error(f"Failed to parse text: {e}")
            raise
    
    def _parse_csv(self, file_content: bytes) -> str:
        """Parse CSV file and extract text"""
        try:
            import csv
            
            # Decode file content
            text_content = self._parse_text(file_content)
            
            # Parse CSV and format as readable text
            csv_reader = csv.reader(io.StringIO(text_content))
            text = ""
            
            for row_num, row in enumerate(csv_reader):
                if row:
                    text += f"Row {row_num + 1}: " + " | ".join(row) + "\n"
            
            return text.strip()
        except Exception as e:
            logger.error(f"Failed to parse CSV: {e}")
            raise
    
    def _parse_image(self, file_content: bytes) -> str:
        """Parse image file and extract text using OCR"""
        try:
            # Open image
            image = Image.open(io.BytesIO(file_content))
            
            # Convert to RGB if necessary
            if image.mode != 'RGB':
                image = image.convert('RGB')
            
            # Extract text using OCR
            text = pytesseract.image_to_string(image, lang='eng+kor')
            
            return text.strip()
        except Exception as e:
            logger.error(f"Failed to parse image: {e}")
            # Return empty string for images that can't be processed
            return ""
    
    def _extract_metadata(self, file_content: bytes, file_info: Dict[str, Any]) -> Dict[str, Any]:
        """Extract additional metadata from file"""
        metadata = {
            'file_size': file_info['size'],
            'mime_type': file_info['mime_type'],
            'extension': file_info['extension']
        }
        
        try:
            if file_info['mime_type'] == 'application/pdf':
                # Extract PDF metadata
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
                if pdf_reader.metadata:
                    metadata.update({
                        'title': pdf_reader.metadata.get('/Title', ''),
                        'author': pdf_reader.metadata.get('/Author', ''),
                        'subject': pdf_reader.metadata.get('/Subject', ''),
                        'creator': pdf_reader.metadata.get('/Creator', ''),
                        'producer': pdf_reader.metadata.get('/Producer', ''),
                        'creation_date': str(pdf_reader.metadata.get('/CreationDate', '')),
                        'modification_date': str(pdf_reader.metadata.get('/ModDate', ''))
                    })
                metadata['page_count'] = len(pdf_reader.pages)
            
            elif file_info['mime_type'] == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                # Extract DOCX metadata
                doc = Document(io.BytesIO(file_content))
                metadata.update({
                    'paragraph_count': len(doc.paragraphs),
                    'table_count': len(doc.tables)
                })
            
            elif file_info['mime_type'] in ['application/vnd.openxmlformats-officedocument.spreadsheetml.sheet', 'application/vnd.ms-excel']:
                # Extract Excel metadata
                workbook = openpyxl.load_workbook(io.BytesIO(file_content))
                metadata.update({
                    'sheet_count': len(workbook.sheetnames),
                    'sheet_names': workbook.sheetnames
                })
            
        except Exception as e:
            logger.warning(f"Failed to extract metadata: {e}")
        
        return metadata
    
    def get_supported_extensions(self) -> List[str]:
        """Get list of supported file extensions"""
        return list(self.extension_mapping.keys())
    
    def get_supported_mime_types(self) -> List[str]:
        """Get list of supported MIME types"""
        return list(self.supported_types.keys())
    
    def is_supported(self, filename: str) -> bool:
        """Check if file type is supported"""
        file_ext = Path(filename).suffix.lower()
        mime_type = self.extension_mapping.get(file_ext, '')
        return mime_type in self.supported_types


# Global file parser instance
_file_parser = None

def get_file_parser() -> FileParser:
    """Get global file parser instance"""
    global _file_parser
    if _file_parser is None:
        _file_parser = FileParser()
    return _file_parser
